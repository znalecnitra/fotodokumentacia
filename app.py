import streamlit as st
from google import genai
from google.genai import types
from PIL import Image, ImageOps
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

# 1. NASTAVENIE STRÁNKY (Široké rozloženie bez bočného panelu)
st.set_page_config(page_title="Znalec - Fotodokumentácia", layout="wide")
st.title("📸 Profesionálna Fotodokumentácia pre posudky")

# 2. VSTUPNÉ ÚDAJE PRE HLAVIČKU A PÄTU
st.subheader("📝 Základné údaje posudku")
col1, col2, col3, col4 = st.columns([2, 2, 4, 3])
with col1:
    znalec = st.text_input("Meno znalca:", value="Ing. Jozef HRČKA")
with col2:
    datum = st.text_input("Dátum obhliadky:", value="08.07.2026")
with col3:
    objekt = st.text_input("Objekt / Byt (Zápätie):", value="Byt č.7, 3.p., o.č.6, Javorová 643/6, k.ú. Chrenová, Nitra")
with col4:
    # API kľúč priamo na hlavnej obrazovke, aby nezavadzal v boku
    gemini_key = st.text_input("🔑 Bezplatný Gemini API kľúč:", type="password")

# Inicializácia AI klienta podľa nového Google štandardu
client = None
if gemini_key:
    try:
        client = genai.Client(api_key=gemini_key)
    except Exception as e:
        st.error(f"Chyba inicializácie AI: {e}")
else:
    st.warning("Upozornenie: Bez vloženého API kľúča nebude fungovať automatické popisovanie fotiek.")

# 3. HROMADNÉ NAHRÁVANIE FOTIEK
st.subheader("🖼️ Nahrať fotky pre projekt")
uploaded_files = st.file_uploader("Vyberte alebo potiahnite fotky naraz:", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

# Inicializácia pamäte poradia
if 'photo_order' not in st.session_state:
    st.session_state.photo_order = []
if 'photo_catalog' not in st.session_state:
    st.session_state.photo_catalog = {}

if uploaded_files:
    # Spracovanie súborov
    for f in uploaded_files:
        if f.name not in st.session_state.photo_catalog:
            # Oprava otočenia fotografie z mobilu
            raw_img = Image.open(f)
            corrected_img = ImageOps.exif_transpose(raw_img)
            
            # Konverzia na bajty
            img_byte_arr = io.BytesIO()
            corrected_img.save(img_byte_arr, format="JPEG")
            final_bytes = img_byte_arr.getvalue()
            
            # Zavolanie nového rozhrania Gemini 2.0 / 2.5
            ai_desc = ""
            if client:
                try:
                    prompt = "Hovoríš po slovensky. Analyzuj tento obrázok zo znaleckej obhliadky nehnuteľnosti. Napíš stručný, profesionálny jednovetný názov toho, čo presne vidíš (napr. 'Pohľad na Kúpeľňu', 'Vstupný portál do bytového domu', 'Pohľad na bytový dom z exteriéru'). Odpovedz LEN týmto popisom, nepíš žiadne reči okolo."
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=[corrected_img, prompt]
                    )
                    ai_desc = response.text.strip()
                except Exception as ai_err:
                    ai_desc = f"Pohľad na objekt"
            
            if not ai_desc:
                ai_desc = f"Pohľad na objekt"
                
            st.session_state.photo_catalog[f.name] = {
                "image": corrected_img,
                "desc": ai_desc,
                "bytes": final_bytes
            }
            if f.name not in st.session_state.photo_order:
                st.session_state.photo_order.append(f.name)

    # Odstránenie odobraných fotiek
    current_names = [f.name for f in uploaded_files]
    st.session_state.photo_order = [name for name in st.session_state.photo_order if name in current_names]

    # 4. VIZUÁLNA KONTROLA - PRESNÁ NAPODOBENINA REGULÁRNEHO WORDU (Mriežka 2 vedľa seba)
    st.markdown("---")
    st.subheader("📄 Vizuálny náhľad dokumentu (Rozloženie ako vo Worde)")
    
    order = st.session_state.photo_order
    
    # Prechádzame po dvojiciach (2 fotky na riadok)
    for i in range(0, len(order), 2):
        grid_col1, grid_col2 = st.columns(2)
        
        # 1. Ľavá fotka v mriežke
        with grid_col1:
            name1 = order[i]
            data1 = st.session_state.photo_catalog[name1]
            st.image(data1["image"], use_container_width=True)
            
            # Textové pole hneď pod fotkou, upravené na šírku
            new_desc1 = st.text_input(f"Popis k fotke vľavo ({name1}):", value=data1["desc"], key=f"txt_{name1}_{i}")
            st.session_state.photo_catalog[name1]["desc"] = new_desc1
            
            # Tlačidlá na zmenu poradia umiestnené diskrétne pod textom
            c_up, c_down = st.columns(2)
            with c_up:
                if i > 0:
                    if st.button("🔼 Posunúť skôr", key=f"up_{name1}_{i}"):
                        st.session_state.photo_order[i], st.session_state.photo_order[i-1] = st.session_state.photo_order[i-1], st.session_state.photo_order[i]
                        st.rerun()
            with c_down:
                if i < len(order) - 1:
                    if st.button("🔽 Posunúť neskôr", key=f"dn_{name1}_{i}"):
                        st.session_state.photo_order[i], st.session_state.photo_order[i+1] = st.session_state.photo_order[i+1], st.session_state.photo_order[i]
                        st.rerun()

        # 2. Pravá fotka v mriežke (ak existuje do páru)
        with grid_col2:
            if i + 1 < len(order):
                name2 = order[i+1]
                data2 = st.session_state.photo_catalog[name2]
                st.image(data2["image"], use_container_width=True)
                
                new_desc2 = st.text_input(f"Popis k fotke vpravo ({name2}):", value=data2["desc"], key=f"txt_{name2}_{i+1}")
                st.session_state.photo_catalog[name2]["desc"] = new_desc2
                
                c_up2, c_down2 = st.columns(2)
                with c_up2:
                    if (i+1) > 0:
                        if st.button("🔼 Posunúť skôr", key=f"up_{name2}_{i+1}"):
                            st.session_state.photo_order[i+1], st.session_state.photo_order[i] = st.session_state.photo_order[i], st.session_state.photo_order[i+1]
                            st.rerun()
                with c_down2:
                    if (i+1) < len(order) - 1:
                        if st.button("🔽 Posunúť neskôr", key=f"dn_{name2}_{i+1}"):
                            st.session_state.photo_order[i+1], st.session_state.photo_order[i+2] = st.session_state.photo_order[i+2], st.session_state.photo_order[i+1]
                            st.rerun()
        st.markdown("<br><hr>", unsafe_allow_html=True)

    # 5. GENERÁTOR STRUKTÚROVANÉHO WORDU (.DOCX)
    def generate_docx():
        doc = docx.Document()
        
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
            
            header = section.header
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.text = f"Znalec: {znalec}\t\tFOTODOKUMENTÁCIA\t\t{datum}"
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.text = objekt
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i in range(0, len(order), 2):
            row_cells_img = table.add_row().cells
            row_cells_txt = table.add_row().cells
            
            # Ľavá strana tabuľky vo Worde
            n1 = order[i]
            item1 = st.session_state.photo_catalog[n1]
            img_s1 = io.BytesIO(item1["bytes"])
            p1 = row_cells_img.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.add_run().add_picture(img_s1, width=Inches(3.1))
            
            t1 = row_cells_txt.paragraphs[0]
            t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t1.add_run(item1["desc"]).font.size = Pt(10)
            
            # Pravá strana tabuľky vo Worde
            if i + 1 < len(order):
                n2 = order[i+1]
                item2 = st.session_state.photo_catalog[n2]
                img_s2 = io.BytesIO(item2["bytes"])
                p2 = row_cells_img.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.add_run().add_picture(img_s2, width=Inches(3.1))
                
                t2 = row_cells_txt.paragraphs[0]
                t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                t2.add_run(item2["desc"]).font.size = Pt(10)
                
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # TLAČIDLO NA STIAHNUTIE WORDU
    st.subheader("💾 Export a uloženie")
    docx_data = generate_docx()
    st.download_button(
        label="📥 Stiahnuť hotovú prílohu vo WORD (.docx)",
        data=docx_data,
        file_name="Fotodokumentacia.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
